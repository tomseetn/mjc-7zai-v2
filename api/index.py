"""MJC 七仔 v2.2 — MJC-style Word doc, server-side parsing"""
import os, json, io, tempfile
from http.server import BaseHTTPRequestHandler
from urllib.parse import quote
import httpx
from anthropic import Anthropic

_anthropic = None
_supabase   = None

def get_anthropic():
    global _anthropic
    if _anthropic is None:
        _anthropic = Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])
    return _anthropic

def get_supabase():
    global _supabase
    if _supabase is None:
        from supabase import create_client
        _supabase = create_client(os.environ["SUPABASE_URL"], os.environ["SUPABASE_ANON_KEY"])
    return _supabase

MAX_HISTORY = 20

MJC_LOGO_B64 = 'iVBORw0KGgoAAAANSUhEUgAAAOEAAADhCAMAAAAJbSJIAAABwlBMVEX////+AAD1+SAAAND0+QAAAMwAAADy9BX1+RX1+Q7///vx8wD/+vr///3/9/f8/cv/7e3///X+/un+LCz5+f7/29v/5ub+b2/9/tf+/uL8/cP/xMT/1dX2+kP/4OD7/bb/y8v+XV3r6/v09PT5+4X+np7/rKz6/KH3+lb+VFT+hYX+IyP+dXX+TEz+vb34+nD+//D7/KxERNYVFdFZWdr+fX3+Jyeuruv/tbXW1vX+RUWLAADFxfHi4vj+k5Nzc97+Ojr4+mf5+475+4j+j4+4uO6SkuWFheI5OdWAgOHh4eGTKSn4+2xnbm7JycmkpOmfAAAgINKWluY/QNVOTthfX9ukYWGXnZ3kJiY5UFCbQ0NfX1pTHBzPOzxreXk8PaHQ0T6fnzJ/fj15eEfNzwA5ODg8PHGNjjURES3f4UW4R0ecpKQbGyLGx0oAACwkODilpFI0NK6wUlI/P4wrKyA5AABmOztFRWFaWgMEBBfkAAC2AAEmJgBBLy8gAABPTqmsrRRzKSkuLjw/P00IPT0AGRnKzBMnJjxfX5BVVSxnZ3RMTEw9PChiYrqamafR0o1tAAC9vclxcVw7OxMAAIkYGAB0dtRKAAANaUlEQVR4nO2c+WPTyBXHLXtkpODYOZzEORzbOWyHkDh34pDDDiEEysIGwoaFbgq7S1tg6W7b7QHddlu6ZVvYq6X9f6sZjTTzRpKvRLZp5/MLZiRZ+ubNvPfmzciBgEQikUgkEolEIpFIJBKJRCKRSCQSiUQikUgkEolEIpFIJBKJRCKRSCSS/2O6BxOJWCyWGJzrbvWjnDbdifmdcjKIbPRgfmforZQZ6ctuRWFTT6qApemaFuTRdISSuZ7WPGaDRIaXS0rvMGiL7SQRgtJ4lQjlh1r0tHUzOjmlKFBfx1BBM7pjZTSUHGzZQ9fOyLlpBeub5dpi1yoYD2rcadmD10bfLpGnTK+xtsGdYG3yCCjfxj4nmu1VTH1Zu61jPlm1c0L0YLtKHDmvUCbttrkdrQ7zWT21LSXa5lOULbtxsFCn+WyJnS2U4sqobT5l2Q6AiXJj+gz0civVOElP2fpmRq3GRL7+7slA860UJMC6p6LYAeJk+rDEuVZq4ujcYvLYAOwpn1Cf0U8LrZTFmOT0ZbpoY/fEifVhI7aDP12bZvqm01ZrrmH/AhXmWqjMZJbTxyJgKohOQx+mleIM0iVO35TVQRNJT32apuv2fLCWXoxirdTXt87pm7amEHMF1wGI534omCxP5OZTBkOp3MVyUq+aiestTMG7Bjh9yjmrOeUyAI05H0pOzCfmxCSlI5Grkq9q+SbLsons8vpKfbS5Jy92UENdsJxLeCdgg7mKgxY1R4+DLK+PeZic0OmMXlierz6bna/QV1FLihrpXl7fej9tHoQeRkfBiVhHTV/Y4+2bWuFqujI1GJDIc14bHysWixvFzTHxwISXRNT8os0k0NdrJdm8ATUUvJhwXrmxpNqEFq5cAjJ3znooTDVBE09fCQi0XeiObUAN6QWvnrW5txiyUdWlDe5Y+eyZdlC4DPQpadrMxpGOkqmKueSeGuJE3tq0D3ScPeMmsbm9dBbqy0Roe4oa0DDfRFXPObYQ4jXesA8MGRKdGpvpabqmoECrztRZQLb5anKdl4EZr9jtSTeJqHmlUxgClRXbxQR1Yj7tYs3PAiVuW80xrFCU2LT5U9cMFHjeOkB6KDZfPWWjQ9BRV63mMy4KtVOV4c0a1KfYlV4cxxAquMSGSoypQKIVNgrObqo1pxoVFUZgyZ7IGz4UBXP1d6SbvMKQNRRzToXNKUaJBly2DiQ03D0b+cpNYMRQ3GxNuShsQloaHRAE2j10HqF8o758H3RTGvljDoXNmDwNC/oUuxZaaFyfEfiBDS+bjT0OhU3IaM4L+qasKN8TzNfpXgCrQOGh2TjoVFjb3KRx+qYFgWwqXz6JvkDgEhiIB2ajQ6HvNYxJQZ89BDtTJ9MnKgyZjQkxNfW55i0GeUWxahWBky8KwV5KFcYEhfrFE9+nEmKMUHq7ql9UMzAgqmZjTpgG+5uxiTFCmTnVr78NFN4yGwuw7OZrwburJAo8X/2ieoDDkCbfSaEgdbq3BKRFfdyS7qlQhApvksZu2En9nBluOQRmq19UFzDgq0XSmAIKfUxnIhmHwHRga9ztzFGXxpo4AAoXzcYy6KT+RYpRMcqTcozy3nXnqQOzR/TT6vaVvb04d2hwvpDP58sT7vVgoZOac4s5YEL/8rVZhz5lBIeOO+87Ts0qd+/Rj9v4Qe9Zf4TOebxzDaMjFNxxioSelHbSHO9J/Vv93XUVGDCC/4+vCqeOKMqP3qGfcXXpVpiekdKwNTQdQwoAZUEjnACH9s1W4GZ0vxJS5xAkiUy/8e8Hx/DULoVTiMugPwmbn0lhSg/ef/DhRx8+uI/Voo+PwDDehiY0qxhDfCdFJ80KPehacbcgsez0Q+hsSpxCYpSfki7bgUunevJn4Z9/d308Ho8bHhI9Cj/krxSmv9TP8MHQr1g/4tSnpMkR8vHxE/5kMq2yFG4Yz7z4CXE7SWM06U/DL5hjSjz6BRAYWKpmQi3pj0AXH6MMc0d6P+VONmuLlsIbeDSFsahrxoNqn4VBj44fA+sL0wqak4JQ6E+gcEyVFHu2RCcZd5ldqLkvUIVXjCf9ZXjctIT+IHzk8v0W8RAUeIm07nAm9GkQLrsIpAtnfYpyBwfJdTtgYC/zK04hdqW//twYhNjhfxw+druBBYwUdHo/yAv0p7zm4kTtZNsYcr95jBt+awWMXkPcXU4hftTfvWNOf7TfP6t0I7GPmrXSPHMzyJ9J4ZSLQGu6FFWU5+GruOUDKsiYWf0h/JgpxK704Isn5mA6U1cfNd0Ml5H6syExsu4iULFqTsYAvXCPzBen/zhOG56H399iCnEStmTE+wQ24cuwmBrwHEKFS6SR66O6L2404pgMYuys2hiChpvsw22P8QjD9cU/hceXmUJckvgyjAuohsI/hyvc6rLQR81UNmj3US3pRy4TdYnzXN13FlssYDrUlU/HSX7zF0Nyhim8QuP9hI4VvvC+1YYg0ExI2Rq+Pzufo71uAtmU3ujBF44DNCjefYJN+hgPtRmm0Oh6iy+OLIV/9byVkI/S1dEY66O+1PCjzskSHnH2cdw7SSzH0pTSV4awqb+9DhB/ailcpPGeKHz5tee9QCnfql10c27Uj0DobkFWN8Sh4rk5skhKcMfIbV7j2E4yOVMhNs2XuG1HJ77U614w4bbKwCwf9ads4S6QvUgQxTqO7I+Y94hJI/jjuyTZxqPr70a8N3NL/R8uc2XMHuyjITMSXrP7qD97ElzDhLLCTsCGs9y/WWB8bMa7LqYQ1z7v4aFKUhrt6SvXW60Kg9DchcGyNX8m9Y6qtskIO8MYfM+/op9JwJghg5D+513SS425wm1zqE6Qafo3bjUddzea8tmCjqqvCVcaHWadNEC69Io5CGnBkSjEacq3ZiupB2qPeCPS+CZUZuiKYcJngefcBSrc65BTXCclUyZzEAZo0Z8oxJnmKzoDNEfi/e+sK+ZyaAJrHAu5CWS1J3+cjNt0CbPLThk1/vucTQsj1iC0riYKjRiw8IXVPE8knvm+Z26uJ5bLI6S/PPYS2G3nMv7MlxyLuy4mTF/44YcwN7VfO35tfVy2FOIB9i2z8xB+iUuzX+7NfxZ+3ekusNOKE5rmy66gfi+By/xZ169ff8L7jav2fyJPHj58+MlxII6LUK8+5875/uUZXGQzHj//9KNnxh9IEBgynQypd5jJtj9LTK6pDKav+rU8t3ByEgYlnPg/j/715s2bb56Fvz42Ru3YIhRIN+vlqUCU96dwmPES2FvX12zgp99/6JbGjFN7C150n1bGy9TLIJ9Kv1kvgfUtMt3CT7/wgp8RZoVFHBgH1du0uWAJ9KluOOopsL5OqoYW//0s/JoNzkm2l8EEFi3sXXpUoObbCprnIGQz+5qIH/3n9ZFtwP5zji5wAwq0NpPSGaEe9GtrpbhJhqPU8Jem8dBeg21Csk2dqGVB5NuuPOcKL2Ogsa+MZEmhAP54QmAbCFywVt8sgf4t01foozAa1kofLbaOwObbYKPsZavZ9KK65t8itlc62qDCyJpVyOoH7XGwnztkb8s3BfoVBTGeyQyh3i0Xo/bfawVutxkDC9n79rsV5otRvm4kcSv+MuraNdM/ycp0U/DQhnsP7SCpmh70aXmQUMnNKHwJqhrRLF8gEDzUTSDQ8qGBbiIQFXz9KQH3wgxjpPpXGERnM+CqSXA0fsgJZHvxAz14uqTr/r4j4rZICMhU/w5RnhglNvkhuMje/Engbf3oms+vF1QzIfd7Ae70Z0V5ygp0onyixr0vYq4tan6/5OM57eVIe14dTZ9z+QtBq8e5MK8eFNkBvPKGCn7v+a3iSCmu04uu4V330iM8u8i/nHaTtXeW8XsL/r/DVDkWck89wiXgkf7hLY+qnCJanPOh6iH3fqHhRJvzy0jOLXmelDIDAwOZzNSM68qUfRYI82O8D73EHegJ6ijflFe0ahdYG3AyyLkYdY8/EEMo2JwXJd02zJyENPh25mLUw03+wDxCuSb9XFAdnbQGMmCyzFyMerAB7jqBJpr2CyUeyxSNAesxV1gPXQUHOpL55r1RH63+2DVTAlG+eMAGYBzcM1Hvi3snou/0BMI81N5/oG4LL9r3NFNfbQlNTUyBGFE8oALV25tet24S3kXS+gCZq52lqQtFrxs3Da/FpvoYAL9FuurhQFvDaSjsTfPfuLlvGlDdbwd9p6IQhoi99tJ3CuPwHOigGwd0/LWLvhrm95XJgBAYN6uh6lLr/QvjRGnpOizgmAsS6nY76TtRTtMLaxsbpoe57PghnVbjusuyBlbggssY7qBq6Gbc4zYtpDFnup4GXxLHOZq6cMn9Fi2mkW4q6CMhvt2GH0fdRhwQCsTFfVUN7bVh97SpXi7lmRTebd5cUtX9S+2sr/Zim0FGWO4MjG2rbdw9bWqM+uui+bCDOVhtb/NRapBYmux3XHbzbTAfZbTyWMxknfICq4c33grzWTh+LcFiBtS6bYo33hrz2Yw4VlhK57Ou6jBvlfkY0ZG1yd3l5eXdrexsur+ufUISiUQikUgkEolEIpFIJBKJRCKRSCQSiUQikUgkEolEIpFIJBKJRCKRSP63+C9lmVwKNF0RLgAAAABJRU5ErkJggg=='
SYSTEM_PROMPT = '你是 MJC 七仔，MJC Leisure Sdn Bhd 专属内部旅游报价 AI 助理。\n\n【角色定位】\n- 服务对象：MJC 内部销售团队\n- 主要任务：收集行程需求，生成行程卡和 Word 报价单\n\n【对话流程】\n每次只问 1-2 个问题，逐步收集以下信息：\n  - 目的地 / 路线\n  - 出行日期（具体日期，如：2026年05月29日）\n  - 人数（报价基准人数，如：30人）\n  - 住宿等级（3星/4星/5星）+ 酒店名称（如有）\n  - 特殊需求（餐数、餐标、门票等）\n  - 交通类型（如：40座旅游巴士）\n  - 2人1房每人价格（如：RMB3150）\n  - 单间差价（如：RMB1380）\n\n信息基本齐全后提示：\n  "✅ 资料基本齐全了！输入「出」生成行程卡，或「word」下载 Word 报价单"\n\n【行程生成触发词：出 / 出行程 / 出单 / 生成 / generate / 出吧 / ok出】\n只输出以下格式，头尾用标记，不要有其他文字：\n\n[ITINERARY_START]\n<div class="iti-wrap">\n  <div class="iti-header">\n    <img src="data:image/png;base64,__LOGO__" class="iti-logo" alt="MJC">\n    <div class="iti-company">\n      <div class="iti-cname">MJC Leisure Sdn Bhd</div>\n      <div class="iti-csub">专业马来西亚入境旅游 · Since 2006</div>\n    </div>\n  </div>\n  <div class="iti-title">[行程标题]</div>\n  <div class="iti-meta">[人数] · [住宿等级] · [日期或待定]</div>\n  <table class="iti-table">\n    <thead><tr><th>日期</th><th>膳食</th><th>行&nbsp;&nbsp;&nbsp;程</th><th>住宿</th></tr></thead>\n    <tbody>\n      [每天：<tr><td><b>第X天</b><br><small>月日</small></td><td class="meals">早<br>中<br>晚</td><td>路线<br>活动</td><td>城市</td></tr>]\n      [膳食：含用汉字 早/中/晚，不含用 X]\n    </tbody>\n  </table>\n  <div class="iti-footer">⚡ 由 MJC 七仔 AI 生成 · 仅供内部参考，请人工审核后再发给客户</div>\n</div>\n[ITINERARY_END]\n\n【Word 报价单触发词：word / Word / 导出 / 下载】\n如果以下字段尚未确认，先问清楚再生成：\n  出行日期、报价人数、2人1房价格、单间差价、酒店全名、交通类型、餐标\n\n确认后只输出以下 JSON，头尾用标记，不要有其他文字：\n\n[WORD_JSON_START]\n{\n  "title": "吉隆坡5天4晚",\n  "date": "2026年05月29日",\n  "pax_for_quote": "30",\n  "price_double": "3150",\n  "price_single_supp": "1380",\n  "hotel": "吉隆坡喜来登福朋China Town 4*酒店或者同级（4晚）",\n  "transport": "40座旅游巴士",\n  "guide": "中文导游",\n  "meals_detail": "酒店含早餐、6次正餐（餐标RMB130）",\n  "tickets": "云顶缆车、马六甲游船",\n  "gratuity": "包含司机导游小费",\n  "insurance": "当地旅行社责任险",\n  "days": [\n    {"day": "第一天", "date": "5月29日", "meals": "X/X/X", "activity": "深圳-吉隆坡\\n依航班时间，送往吉隆坡酒店！晚上亚罗街自理", "stay": "吉隆坡"},\n    {"day": "第二天", "date": "5月30日", "meals": "早/中/晚", "activity": "吉隆坡-马六甲-吉隆坡\\n上午：鸡场街、荷兰红屋、马六甲游船\\n下午：返回吉隆坡", "stay": "吉隆坡"}\n  ]\n}\n[WORD_JSON_END]\n\nmeals 字段规则：早/中/晚 格式，含→用汉字（早中晚），不含→用 X\n  全不含：X/X/X | 仅含早：早/X/X | 全含：早/中/晚\n\n【语言】默认中文  【安全】所有行程需经 MJC 同事审核后发给客户'
HTML_PAGE = '<!DOCTYPE html>\n<html lang="zh">\n<head>\n<meta charset="UTF-8">\n<meta name="viewport" content="width=device-width, initial-scale=1.0">\n<title>MJC 七仔 Demo</title>\n<style>\n*{box-sizing:border-box;margin:0;padding:0}\nbody{font-family:-apple-system,BlinkMacSystemFont,"Segoe UI",sans-serif;background:#f0f2f5;height:100vh;display:flex;flex-direction:column;align-items:center;justify-content:center}\n.chat-container{width:100%;max-width:560px;height:93vh;background:#fff;border-radius:16px;box-shadow:0 4px 24px rgba(0,0,0,.12);display:flex;flex-direction:column;overflow:hidden}\n.header{background:#075E54;color:#fff;padding:14px 18px;display:flex;align-items:center;gap:12px;flex-shrink:0}\n.avatar{width:40px;height:40px;background:#25D366;border-radius:50%;display:flex;align-items:center;justify-content:center;font-size:20px}\n.header-info h3{font-size:16px;font-weight:600}\n.header-info p{font-size:12px;opacity:.8}\n.messages{flex:1;overflow-y:auto;padding:12px;display:flex;flex-direction:column;gap:6px;background:#E5DDD5}\n.msg{max-width:80%;padding:8px 12px;border-radius:8px;font-size:14px;line-height:1.55;word-wrap:break-word}\n.msg.bot{background:#fff;align-self:flex-start;border-top-left-radius:2px;box-shadow:0 1px 2px rgba(0,0,0,.1)}\n.msg.user{background:#DCF8C6;align-self:flex-end;border-top-right-radius:2px;box-shadow:0 1px 2px rgba(0,0,0,.1)}\n.msg.typing{background:#fff;align-self:flex-start;color:#999;font-style:italic}\n.msg.itinerary{max-width:100%;padding:0;background:transparent;align-self:stretch;box-shadow:none}\n.iti-wrap{background:#fff;border-radius:10px;overflow:hidden;box-shadow:0 2px 12px rgba(0,0,0,.15);border:1px solid #ddd}\n.iti-header{background:linear-gradient(135deg,#075E54,#128C7E);padding:16px 20px;display:flex;align-items:center;gap:14px}\n.iti-logo{width:58px;height:58px;object-fit:contain;background:#fff;border-radius:50%;padding:4px}\n.iti-cname{color:#fff;font-size:16px;font-weight:700}\n.iti-csub{color:rgba(255,255,255,.8);font-size:11px;margin-top:3px}\n.iti-title{background:#f8f8f8;padding:12px 16px;font-size:16px;font-weight:700;color:#075E54;border-bottom:2px solid #075E54}\n.iti-meta{padding:7px 16px;font-size:12px;color:#666;background:#fafafa;border-bottom:1px solid #eee}\n.iti-table{width:100%;border-collapse:collapse;font-size:13px}\n.iti-table th{background:#075E54;color:#fff;padding:8px 10px;text-align:left;font-weight:600;font-size:12px}\n.iti-table td{padding:8px 10px;border-bottom:1px solid #eee;vertical-align:top}\n.iti-table tr:nth-child(even) td{background:#f9fffe}\n.iti-table td.meals{font-size:13px;color:#444;white-space:nowrap;line-height:2;text-align:center}\n.iti-footer{background:#f0f2f5;padding:8px 16px;font-size:11px;color:#999;text-align:center;border-top:1px solid #eee}\n.tip-bar{background:#e8f5e9;border-top:1px solid #c8e6c9;padding:6px 14px;font-size:12px;color:#2e7d32;text-align:center;flex-shrink:0}\n.input-area{padding:10px 14px;background:#f0f2f5;display:flex;gap:8px;align-items:flex-end;flex-shrink:0}\ntextarea{flex:1;padding:10px 14px;border:none;border-radius:20px;outline:none;font-size:14px;font-family:inherit;resize:none;max-height:100px;background:#fff;box-shadow:0 1px 3px rgba(0,0,0,.1)}\nbutton{width:42px;height:42px;background:#075E54;color:#fff;border:none;border-radius:50%;cursor:pointer;font-size:18px;display:flex;align-items:center;justify-content:center;flex-shrink:0;transition:background .2s}\nbutton:hover{background:#128C7E}\nbutton:disabled{background:#ccc;cursor:not-allowed}\n.demo-badge{text-align:center;padding:4px;font-size:11px;color:#999;background:#f0f2f5;flex-shrink:0}\n</style>\n</head>\n<body>\n<div class="chat-container">\n  <div class="header">\n    <div class="avatar">🤖</div>\n    <div class="header-info">\n      <h3>MJC 七仔</h3>\n      <p>内部旅游报价助理 · Demo 版</p>\n    </div>\n  </div>\n  <div class="messages" id="messages">\n    <div class="msg bot">你好！我是 MJC 七仔 👋<br>请问这次行程的目的地和大概天数是什么？</div>\n  </div>\n  <div class="tip-bar">💡 输入「<b>出</b>」生成行程卡 · 输入「<b>word</b>」下载 Word 报价单</div>\n  <div class="demo-badge">⚡ Demo 演示版 — Powered by Claude AI + FHA</div>\n  <div class="input-area">\n    <textarea id="inp" placeholder="输入消息..." rows="1"></textarea>\n    <button id="sendBtn">➤</button>\n  </div>\n</div>\n<script>\nvar chatHistory = [];\n\nfunction sendMsg() {\n  var inp = document.getElementById(\'inp\');\n  var msg = inp.value.trim();\n  if (!msg) return;\n  var wordTrigger = /^(word|出word|出 word|导出|下载)$/i.test(msg);\n  addMsg(msg, \'user\');\n  inp.value = \'\';\n  inp.style.height = \'auto\';\n  document.getElementById(\'sendBtn\').disabled = true;\n\n  if (wordTrigger) {\n    var typing = addMsg(\'正在生成 Word 报价单...\', \'typing\');\n    fetch(\'/api/word\', {\n      method: \'POST\',\n      headers: {\'Content-Type\': \'application/json\'},\n      body: JSON.stringify({messages: chatHistory})\n    }).then(function(r) {\n      typing.remove();\n      if (!r.ok) return r.json().then(function(e){ throw new Error(e.error || \'error\'); });\n      var disp = r.headers.get(\'Content-Disposition\') || \'\';\n      var m = disp.match(/filename\\*?=(?:UTF-8\'\')?(.+)/i);\n      var fname = m ? decodeURIComponent(m[1].replace(/"/g,\'\')) : \'MJC行程单.docx\';\n      return r.blob().then(function(blob){ return {blob: blob, fname: fname}; });\n    }).then(function(obj) {\n      var url = URL.createObjectURL(obj.blob);\n      var a = document.createElement(\'a\');\n      a.href = url; a.download = obj.fname;\n      document.body.appendChild(a); a.click();\n      document.body.removeChild(a); URL.revokeObjectURL(url);\n      addMsg(\'✅ Word 报价单已下载：\' + obj.fname, \'bot\');\n      document.getElementById(\'sendBtn\').disabled = false;\n    }).catch(function(e) {\n      addMsg(\'❌ 生成失败：\' + e.message, \'bot\');\n      document.getElementById(\'sendBtn\').disabled = false;\n    });\n    return;\n  }\n\n  chatHistory.push({role: \'user\', content: msg});\n  var typing = addMsg(\'七仔正在思考...\', \'typing\');\n  fetch(\'/api/chat\', {\n    method: \'POST\',\n    headers: {\'Content-Type\': \'application/json\'},\n    body: JSON.stringify({messages: chatHistory})\n  }).then(function(r){ return r.json(); }).then(function(data) {\n    typing.remove();\n    if (data.type === \'itinerary\' && data.html) {\n      var div = document.createElement(\'div\');\n      div.className = \'msg itinerary\';\n      div.innerHTML = data.html;\n      document.getElementById(\'messages\').appendChild(div);\n      div.scrollIntoView({behavior: \'smooth\'});\n      chatHistory.push({role: \'assistant\', content: \'[行程卡已生成]\'});\n    } else {\n      var reply = data.reply || \'抱歉，出错了，请重试\';\n      chatHistory.push({role: \'assistant\', content: reply});\n      addMsg(reply, \'bot\');\n    }\n    document.getElementById(\'sendBtn\').disabled = false;\n  }).catch(function(e) {\n    typing.remove();\n    addMsg(\'连接错误，请检查网络\', \'bot\');\n    document.getElementById(\'sendBtn\').disabled = false;\n  });\n}\n\nfunction addMsg(text, type) {\n  var div = document.createElement(\'div\');\n  div.className = \'msg \' + type;\n  if (type === \'user\') {\n    div.textContent = text;\n  } else {\n    div.innerHTML = text.replace(/&/g,\'&amp;\').replace(/</g,\'&lt;\').replace(/>/g,\'&gt;\')\n      .replace(/\\n/g,\'<br>\').replace(/\\*\\*([^*]+)\\*\\*/g,\'<b>$1</b>\');\n  }\n  document.getElementById(\'messages\').appendChild(div);\n  div.scrollIntoView({behavior: \'smooth\'});\n  return div;\n}\ndocument.getElementById(\'sendBtn\').addEventListener(\'click\', sendMsg);\ndocument.getElementById(\'inp\').addEventListener(\'input\', function() {\n  this.style.height = \'auto\';\n  this.style.height = Math.min(this.scrollHeight, 100) + \'px\';\n});\ndocument.getElementById(\'inp\').addEventListener(\'keydown\', function(e) {\n  if (e.key === \'Enter\' && !e.shiftKey) { e.preventDefault(); sendMsg(); }\n});\n</script>\n</body>\n</html>'
STD_EXCLUSIONS = '马来西亚签证、公务自理、旅游意外险、机票与机场税、酒店行李搬运、自费项目与个人消费，行程中未提到的其他费用，行程以外的景点，医疗费，交通延阻，罢工及其他不可抗拒的因素所致的额外费用。'
STD_REMARK = '备注：酒店价格会根据酒店的房态而更改，我处并没有预定任何房间。'

WORD_GEN_PROMPT = '你是行程单数据提取助手。根据对话历史，提取行程信息，输出以下格式。\n\n只输出 JSON，不要有任何其他文字：\n\n[WORD_JSON_START]\n{\n  "title": "目的地+天数（如：吉隆坡5天4晚）",\n  "date": "出行日期（如：2026年05月29日）",\n  "pax_for_quote": "报价人数（纯数字，如：30）",\n  "price_double": "2人1房每人价格（纯数字，如：3150）",\n  "price_single_supp": "单间差价（纯数字，如：1380）",\n  "hotel": "酒店名称+星级+晚数（如：吉隆坡喜来登福朋4*酒店或同级（4晚））",\n  "transport": "交通类型（如：40座旅游巴士）",\n  "guide": "中文导游",\n  "meals_detail": "用餐详情（如：酒店含早餐、6次正餐（餐标RMB100））",\n  "tickets": "门票项目（如：云顶缆车、马六甲游船）",\n  "gratuity": "包含司机导游小费",\n  "insurance": "当地旅行社责任险",\n  "days": [\n    {"day": "第一天", "date": "5月29日", "meals": "X/X/X", "activity": "出发城市-吉隆坡\\n依航班时间，抵达后入住酒店", "stay": "吉隆坡"},\n    {"day": "第二天", "date": "5月30日", "meals": "早/中/晚", "activity": "吉隆坡市区\\n景点活动描述", "stay": "吉隆坡"}\n  ]\n}\n[WORD_JSON_END]\n\nmeals 规则：含→汉字（早/中/晚），不含→X。第一天通常全X，最后一天早/X/X。\n如信息不全，用合理默认值填入。行程天数根据对话中的天数生成对应行数。'

# ─── helpers ────────────────────────────────────────────────

def get_history(phone):
    try:
        res = (get_supabase().table("conversations")
               .select("role,content").eq("phone_number", phone)
               .order("created_at", desc=True).limit(MAX_HISTORY).execute())
        rows = (res.data or [])[::-1]
        return [{"role": r["role"], "content": r["content"]} for r in rows]
    except Exception as e:
        print(f"[Supabase] {e}")
        return []

def save_message(phone, role, content):
    try:
        get_supabase().table("conversations").insert(
            {"phone_number": phone, "role": role, "content": content}).execute()
    except Exception as e:
        print(f"[Supabase] {e}")

def send_wati(phone, msg):
    inst = os.environ.get("WATI_INSTANCE_URL","").rstrip("/")
    tok  = os.environ.get("WATI_API_TOKEN","")
    if not inst or not tok: return
    try:
        httpx.post(f"{inst}/api/v1/sendSessionMessage/{phone}",
                   json={"messageText": msg},
                   headers={"Authorization": f"Bearer {tok}"},
                   timeout=15).raise_for_status()
    except Exception as e:
        print(f"[WATI] {e}")

def ask_claude(messages, system):
    try:
        resp = get_anthropic().messages.create(
            model="claude-haiku-4-5-20251001", max_tokens=4096,
            system=system, messages=messages[-20:])
        return resp.content[0].text
    except Exception as e:
        print(f"[Claude] {e}")
        return "抱歉，七仔暂时无法处理，请稍后再试。"

def parse_block(reply, start_tag, end_tag):
    si = reply.find(start_tag)
    ei = reply.find(end_tag)
    if si != -1 and ei != -1:
        return reply[si+len(start_tag):ei].strip()
    return None

# ─── Word doc builder ────────────────────────────────────────

def build_word_doc(data):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import base64, os as _os

    doc = Document()

    # — Page: A4 portrait, tight margins —
    sec = doc.sections[0]
    sec.page_width  = Cm(21);  sec.page_height = Cm(29.7)
    sec.top_margin  = Cm(1.2); sec.bottom_margin = Cm(1.0)
    sec.left_margin = Cm(1.5); sec.right_margin  = Cm(1.5)

    def shd(cell, hex6):
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        s = OxmlElement("w:shd")
        s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto")
        s.set(qn("w:fill"), hex6); tcPr.append(s)

    def no_border(table):
        tblPr = table._tbl.tblPr
        tb = OxmlElement("w:tblBorders")
        for n in ["top","left","bottom","right","insideH","insideV"]:
            b = OxmlElement(f"w:{n}"); b.set(qn("w:val"),"none"); tb.append(b)
        tblPr.append(tb)

    def sp(para, before=0, after=0):
        para.paragraph_format.space_before = Pt(before)
        para.paragraph_format.space_after  = Pt(after)

    def fnt(run, size, bold=False, color=None):
        run.font.size = Pt(size)
        if bold: run.bold = True
        if color: run.font.color.rgb = RGBColor(*color)

    # — Header: logo left | company right —
    htbl = doc.add_table(1, 2); htbl.alignment = 1; no_border(htbl)
    logo_c = htbl.rows[0].cells[0]
    try:
        lb = base64.b64decode(MJC_LOGO_B64)
        tf = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        tf.write(lb); tf.close()
        logo_c.paragraphs[0].add_run().add_picture(tf.name, width=Cm(1.8))
        _os.unlink(tf.name)
    except Exception: logo_c.paragraphs[0].text = "MJC"
    logo_c.width = Cm(2.2)
    sp(logo_c.paragraphs[0], 2, 0)

    txt_c = htbl.rows[0].cells[1]; txt_c.width = Cm(15.8)
    p1 = txt_c.paragraphs[0]; sp(p1, 3, 0)
    fnt(p1.add_run("MJC Leisure Sdn Bhd"), 11, bold=True, color=(7,94,84))

    p2 = txt_c.add_paragraph(); sp(p2, 0, 0)
    fnt(p2.add_run("专业马来西亚入境旅游 · Since 2006"), 8.5, color=(100,100,100))

    # — Title —
    title_str = data.get("title","马来西亚行程")
    tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER; sp(tp, 5, 3)
    fnt(tp.add_run(title_str), 14, bold=True, color=(7,94,84))

    # — Itinerary table —
    days = data.get("days", [])
    tbl = doc.add_table(1 + len(days), 4); tbl.style = "Table Grid"

    # Fixed layout
    tblL = OxmlElement("w:tblLayout"); tblL.set(qn("w:type"),"fixed")
    tbl._tbl.tblPr.append(tblL)

    # Column widths (total 18cm)
    cw_cm = [2.3, 1.4, 12.1, 2.2]
    for ci, w in enumerate(cw_cm):
        for cell in tbl.column_cells(ci):
            cell.width = Cm(w)
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(int(w*567)))
            tcW.set(qn("w:type"),"dxa")
            tcPr.append(tcW)

    # Header row
    for i, h in enumerate(["日期","膳食","行　程","住宿"]):
        c = tbl.rows[0].cells[i]; c.text = h
        shd(c,"075E54")
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; sp(p,1,1)
        fnt(p.runs[0], 9, bold=True, color=(255,255,255))

    # Data rows
    for ri, day in enumerate(days):
        row = tbl.rows[ri+1]

        # 日期 col
        dc = row.cells[0]; dc.text = ""
        p = dc.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; sp(p,2,0)
        fnt(p.add_run(day.get("day","")), 9, bold=True)
        if day.get("date"):
            p2 = dc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER; sp(p2,0,2)
            fnt(p2.add_run(day["date"]), 8)

        # 膳食 col  — 早/中/晚 format
        mc = row.cells[1]; mc.text = ""
        raw_meals = day.get("meals","X/X/X")
        parts = [x.strip() for x in raw_meals.split("/")]
        while len(parts) < 3: parts.append("X")
        first = True
        for ml in parts[:3]:
            if first:
                mp = mc.paragraphs[0]; first = False
            else:
                mp = mc.add_paragraph()
            mp.text = ml; mp.alignment = WD_ALIGN_PARAGRAPH.CENTER; sp(mp,0,0)
            if mp.runs: fnt(mp.runs[0], 9)

        # 行程 col
        ac = row.cells[2]; ac.text = ""
        lines = day.get("activity","").split("\n")
        first = True
        for ln in lines:
            ln = ln.strip()
            if not ln: continue
            if first:
                ap = ac.paragraphs[0]; ap.text = ln; first = False
            else:
                ap = ac.add_paragraph(ln)
            sp(ap, 0, 0)
            if ap.runs: fnt(ap.runs[0], 9)

        # 住宿 col
        sc = row.cells[3]; sc.text = day.get("stay","")
        sc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER; sp(sc.paragraphs[0],2,2)
        if sc.paragraphs[0].runs: fnt(sc.paragraphs[0].runs[0], 9)

        if ri % 2 == 1:
            for c in row.cells: shd(c,"F0FFFD")

    # — Disclaimer —
    dp = doc.add_paragraph(); sp(dp,5,0)
    fnt(dp.add_run("                     以上行程仅供参考，我社有权根马来西亚假期的情况调整！"), 8.5, bold=True)

    # — Pricing —
    pax  = data.get("pax_for_quote","")
    pd   = data.get("price_double","")
    ps   = data.get("price_single_supp","")
    if pax:
        pp = doc.add_paragraph(); sp(pp,3,0)
        fnt(pp.add_run(f"报价以{pax}人为准"), 9, bold=True)
    if pd:
        p2 = doc.add_paragraph(); sp(p2,0,0)
        fnt(p2.add_run(f"2人1房 每人RMB{pd}"), 9)
    if ps:
        p3 = doc.add_paragraph(); sp(p3,0,3)
        fnt(p3.add_run(f"1人1房补单间差房 RMB{ps}"), 9)

    # — Details 1-8 —
    fields = [
        ("日期", data.get("date","")),
        ("酒店", data.get("hotel","")),
        ("交通", data.get("transport","")),
        ("导游", data.get("guide","")),
        ("用餐", data.get("meals_detail","")),
        ("门票", data.get("tickets","")),
        ("小费", data.get("gratuity","")),
        ("保险", data.get("insurance","")),
    ]
    for i,(lbl,val) in enumerate(fields,1):
        if not val: continue
        fp = doc.add_paragraph(); sp(fp,0,0)
        fnt(fp.add_run(f"{i}、	{lbl}：{val}"), 9)

    # — Exclusions —
    ep = doc.add_paragraph(); sp(ep,4,0)
    fnt(ep.add_run("报价不含："), 9, bold=True)
    etp = doc.add_paragraph(STD_EXCLUSIONS); sp(etp,0,0)
    if etp.runs: fnt(etp.runs[0], 8.5)

    # — Remarks —
    remarks = data.get("remarks", "")
    remark_text = f"备注：{remarks}" if remarks else STD_REMARK
    rp = doc.add_paragraph(remark_text); sp(rp,3,0)
    if rp.runs: fnt(rp.runs[0], 8.5)

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.read(), title_str


# ─── HTTP handler ────────────────────────────────────────────

class handler(BaseHTTPRequestHandler):
    def log_message(self, fmt, *args): print(f"[HTTP] {fmt%args}")

    def do_GET(self):
        path = self.path.split("?")[0]
        if path in ("/","/demo"):
            body = HTML_PAGE.encode("utf-8")
            self.send_response(200)
            self.send_header("Content-Type","text/html; charset=utf-8")
            self.end_headers(); self.wfile.write(body)
        else:
            self._json(200,{"status":"ok","bot":"MJC 七仔 v2.2"})

    def do_POST(self):
        path = self.path.split("?")[0]
        raw  = self.rfile.read(int(self.headers.get("Content-Length",0)))
        try: data = json.loads(raw)
        except: self._json(400,{"error":"invalid JSON"}); return

        if path == "/api/chat":
            msgs = data.get("messages",[])
            if not msgs: self._json(400,{"error":"no messages"}); return
            reply = ask_claude(msgs, SYSTEM_PROMPT)
            html  = parse_block(reply,"[ITINERARY_START]","[ITINERARY_END]")
            if html:
                html = html.replace("__LOGO__", MJC_LOGO_B64)
                self._json(200,{"type":"itinerary","html":html})
            else:
                self._json(200,{"type":"text","reply":reply})

        elif path == "/api/word":
            msgs = data.get("messages",[])
            reply = ask_claude(msgs, WORD_GEN_PROMPT)
            raw_json = parse_block(reply,"[WORD_JSON_START]","[WORD_JSON_END]")
            if not raw_json:
                self._json(500,{"error":"无法生成行程数据，请先和七仔确认行程信息"}); return
            try:
                iti = json.loads(raw_json)
                docx_bytes, title = build_word_doc(iti)
                filename = f"MJC 行程单 2026 {title}.docx"
                self.send_response(200)
                self.send_header("Content-Type","application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                self.send_header("Content-Disposition", f"attachment; filename*=UTF-8\'\'{quote(filename)}")
                self.send_header("Content-Length", str(len(docx_bytes)))
                self.end_headers(); self.wfile.write(docx_bytes)
            except Exception as e:
                print(f"[Word] {e}")
                self._json(500,{"error":str(e)})

        elif path in ("/api/webhook","/webhook"):
            phone   = data.get("waId") or data.get("phone","")
            msg_obj = data.get("text") or {}
            umsg    = (msg_obj.get("body") or data.get("body") or "").strip()
            if not phone or not umsg:
                self._json(200,{"ok":True,"skip":"no phone or body"}); return
            hist  = get_history(phone)
            reply = ask_claude(hist+[{"role":"user","content":umsg}], SYSTEM_PROMPT)
            save_message(phone,"user",umsg); save_message(phone,"assistant",reply)
            send_wati(phone, reply)
            self._json(200,{"ok":True})
        else:
            self._json(404,{"error":"not found"})

    def _json(self, code, body):
        self.send_response(code)
        self.send_header("Content-Type","application/json")
        self.end_headers()
        self.wfile.write(json.dumps(body, ensure_ascii=False).encode())
